"""Generate formal B&W Word progress report for completed EdVidura LTI Hello work."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "EdVidura_Progress_Report_Completed_Work.docx"


def set_run_font(run, size: float = 12, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(
    doc: Document,
    text: str,
    *,
    size: float = 12,
    bold: bool = False,
    center: bool = False,
    space_after: float = 8,
    justify: bool = False,
):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_heading_numbered(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text.upper())
    set_run_font(run, size=13, bold=True)
    return p


def add_subhead(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True)
    return p


def shade_cell(cell, color: str = "E8E8E8") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_border(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:color"), "000000")
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_table(doc: Document, headers: list[str], rows: list[tuple], col_widths: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10.5, bold=True)
        shade_cell(cell, "E8E8E8")
        set_cell_border(cell)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10.5, bold=False)
            set_cell_border(cell)
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Inches(w)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str], ordered: bool = False) -> None:
    style = "List Number" if ordered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        # replace default run text cleanly
        if p.runs:
            p.runs[0].text = item
            set_run_font(p.runs[0], size=12)
        else:
            run = p.add_run(item)
            set_run_font(run, size=12)
        p.paragraph_format.space_after = Pt(3)


def main() -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    # Cover box
    cover = doc.add_table(rows=1, cols=1)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = cover.rows[0].cells[0]
    set_cell_border(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run("EDVIDURA · ENGINEERING PROGRESS REPORT")
    set_run_font(run, size=11, bold=True)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("LTI Hello / Multi-Tenant Spike")
    set_run_font(run, size=18, bold=True)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Completed Work to Date")
    set_run_font(run, size=16, bold=True)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Formal status report of delivered capability through Slice A")
    set_run_font(run, size=12)

    doc.add_paragraph()
    add_table(
        doc,
        ["Field", "Detail"],
        [
            ("Document type", "Formal progress report (black & white / print-ready)"),
            ("Project", "EdVidura LTI Hello (FastAPI + Moodle LTI 1.3 + PostgreSQL)"),
            ("Repository", "https://github.com/meghanachinucom/edvidura-lti-hello"),
            ("Report date", "11 August 2026"),
            (
                "Scope of this report",
                "Work completed and verified in the development spike; excludes deferred modules (AI, XR, federation, enclave)",
            ),
            ("Classification", "Internal — technical / project governance"),
        ],
        col_widths=[2.2, 4.5],
    )

    add_heading_numbered(doc, "1. Executive summary")
    add_para(
        doc,
        "EdVidura LTI Hello is a working demonstration of an LMS-embedded learning tool. "
        "Learners and instructors open the tool from Moodle using LTI 1.3. The application resolves "
        "the correct institution (tenant) from verified platform registration data, stores activity in a "
        "shared PostgreSQL database protected by Row Level Security (RLS), and delivers a vertical product "
        "slice: a short quiz with optional grade return to the Moodle gradebook via Assignment and Grade "
        "Services (AGS).",
        justify=True,
    )
    add_para(
        doc,
        "As of this report, the core launch path, tenancy model, onboarding APIs, tenant-isolation tests, "
        "and Slice A quiz journey are complete. Remaining work is operational hardening, configuration "
        "checklists for demos, and later product slices—not foundational redesign.",
        justify=True,
    )

    add_heading_numbered(doc, "2. Purpose and objectives")
    add_subhead(doc, "2.1 Purpose")
    add_para(
        doc,
        "To prove that EdVidura can run inside a customer LMS without replacing Moodle, while enforcing "
        "strict multi-tenant data separation and returning scores to the institutional gradebook.",
        justify=True,
    )
    add_subhead(doc, "2.2 Objectives achieved")
    add_bullets(
        doc,
        [
            "Secure LTI 1.3 login and launch against a local Moodle instance.",
            "Tenant resolution derived only from verified issuer, client identifier, and deployment.",
            "Shared-database tenancy with PostgreSQL RLS (decision DEC-006).",
            "Institution and student onboarding HTTP APIs.",
            "End-to-end Slice A: launch → three-question quiz → persisted attempt → teacher list → AGS passback.",
            "Automated tests proving cross-tenant isolation on launch events.",
        ],
        ordered=True,
    )

    add_heading_numbered(doc, "3. Technical environment")
    add_table(
        doc,
        ["Layer", "Technology", "Notes"],
        [
            ("Application", "Python 3.11+ / FastAPI / PyLTI1p3", "Runs locally via uvicorn (typically port 8000)"),
            (
                "Database",
                "PostgreSQL 16 (Docker: postgres:16-alpine)",
                "Host port 5433; no native Postgres install required",
            ),
            ("LMS", "Moodle (Docker Compose)", "Host port 8085"),
            ("Tenancy", "Shared schema + RLS", "App role edvidura_app (non-superuser)"),
            ("Grades", "LTI AGS", "Moodle gradebook is system of record"),
        ],
        col_widths=[1.4, 2.6, 2.7],
    )
    add_para(
        doc,
        "Note on database hosting. PostgreSQL is not installed as a Windows application. Docker Desktop "
        "runs the database container; the FastAPI process connects to 127.0.0.1:5433 through Docker port mapping.",
        justify=True,
    )

    add_heading_numbered(doc, "4. Architecture overview")
    arch = (
        "Learner / Instructor\n"
        "        |\n"
        "        v\n"
        "   Moodle (LTI 1.3 platform)\n"
        "        |\n"
        "        |  OIDC login + launch (iss, client_id, deployment)\n"
        "        v\n"
        "   EdVidura FastAPI tool\n"
        "        |-- resolve tenant from lti_platforms (fail closed)\n"
        "        |-- SET LOCAL app.tenant_id  (RLS)\n"
        "        |-- quiz / APIs / teacher views\n"
        "        |\n"
        "        +--> PostgreSQL (shared DB, RLS policies)\n"
        "        |\n"
        "        +--> AGS score POST --> Moodle gradebook (when enabled)"
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(arch)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)

    add_heading_numbered(doc, "5. Completed work (by workstream)")

    add_subhead(doc, "5.1 LTI 1.3 integration")
    add_table(
        doc,
        ["Item", "Status", "Evidence"],
        [
            ("OIDC login endpoint", "Completed", "/lti/login"),
            ("Launch endpoint", "Completed", "/lti/launch → redirect to quiz"),
            ("JWKS publication", "Completed", "/.well-known/jwks.json"),
            ("Unknown platform rejection", "Completed", "Fail-closed on unregistered issuer/client"),
            ("Launch event audit write", "Completed", "launch_events under RLS"),
        ],
        col_widths=[2.4, 1.2, 3.1],
    )

    add_subhead(doc, "5.2 Multi-tenancy and data isolation")
    add_table(
        doc,
        ["Item", "Status", "Evidence"],
        [
            ("Tenant + platform schema", "Completed", "tenants, lti_platforms in db/init.sql"),
            ("Resolution contract", "Completed", "docs/TENANT_RESOLUTION.md"),
            ("RLS on launch events", "Completed", "FORCE RLS + tenant policy"),
            ("RLS on quiz attempts", "Completed", "Tenant-scoped insert/select"),
            ("Isolation CI tests", "Completed", "PR #2 merged; tests/test_tenant_isolation.py"),
            ("Live cross-check endpoint", "Completed", "GET /dev/tenancy/cross-check"),
        ],
        col_widths=[2.4, 1.2, 3.1],
    )

    add_subhead(doc, "5.3 Onboarding APIs")
    add_table(
        doc,
        ["Item", "Status", "Evidence"],
        [
            ("Institution create/list", "Completed", "POST/GET /api/v1/institutions"),
            ("Student create/list", "Completed", "POST/GET /api/v1/students"),
            ("OpenAPI / Swagger", "Completed", "http://127.0.0.1:8000/docs"),
        ],
        col_widths=[2.4, 1.2, 3.1],
    )

    add_subhead(doc, "5.4 Slice A — quiz journey")
    add_table(
        doc,
        ["Item", "Status", "Evidence"],
        [
            ("Three-question quiz UI", "Completed", "/quiz after LTI launch"),
            ("Score calculation", "Completed", "app/quiz_content.py; unit tests"),
            ("Persist attempts", "Completed", "quiz_attempts table + migration"),
            ("Result page", "Completed", "/quiz/result/{id}"),
            ("Teacher attempts list", "Completed", "/teacher/attempts (tenant-scoped)"),
            ("AGS grade passback", "Completed*", "app/ags_passback.py; requires Moodle AGS config"),
            ("Submit UX (non-blocking)", "Completed", "Save + redirect; AGS in background task"),
            ("Launch snapshot persistence", "Completed", "lti_launch_snapshots + quiz session tokens"),
        ],
        col_widths=[2.4, 1.2, 3.1],
    )
    add_para(
        doc,
        "*AGS passback is implemented in code. Successful delivery to Moodle depends on tool/activity "
        "configuration (Assignment and Grade Services enabled; activity set to accept grades).",
        justify=True,
    )

    add_subhead(doc, "5.5 Platform operations (local)")
    add_table(
        doc,
        ["Item", "Status", "Evidence"],
        [
            ("Postgres Compose stack", "Completed", "db/docker-compose.yml"),
            ("Moodle Compose stack", "Completed", "moodle/docker-compose.yml"),
            ("Health endpoint", "Completed", "GET /health (includes DB check)"),
            ("Seed / setup scripts", "Completed", "Key generation; platform seed scripts"),
            ("Project README", "Completed", "Setup and Slice A test steps"),
        ],
        col_widths=[2.4, 1.2, 3.1],
    )

    add_heading_numbered(doc, "6. Key design decisions (locked)")
    add_table(
        doc,
        ["Decision", "Choice", "Rationale"],
        [
            ("LMS for Release 1", "Moodle only", "Fastest path to a demoable LTI product"),
            ("Launch protocol", "LTI 1.3 (PyLTI1p3), new window", "Standard Advantage path; reduces iframe cookie issues"),
            ("Cloud tenancy", "Shared DB + shared schema + RLS", "Simple SaaS spike; isolation enforced in database"),
            ("Regulated / defence", "Dedicated instance later", "Not schema-per-tenant in this spike"),
            ("Official grades", "Moodle gradebook via AGS", "Institution remains system of record"),
            ("Delivery style", "Thin vertical slices", "Prove one journey end-to-end before expanding"),
        ],
        col_widths=[2.0, 2.3, 2.4],
    )

    add_heading_numbered(doc, "7. Verification and acceptance evidence")
    add_bullets(
        doc,
        [
            "Functional: Student launch from Moodle lands on quiz; submit stores attempt; instructor can view tenant-scoped attempts.",
            "Security / tenancy: Isolation tests and /dev/tenancy/cross-check show Tenant A cannot read Tenant B launch data under RLS.",
            "Integration: When Moodle AGS is enabled, score is posted to the gradebook; quiz UI reports AGS availability on launch.",
            "Resilience (dev): Launch JWT and quiz session token persist in PostgreSQL so grade passback survives local server reload.",
        ],
        ordered=True,
    )

    add_heading_numbered(doc, "8. How to inspect stored data")
    add_para(doc, "Data may be reviewed without a desktop PostgreSQL installation:", justify=True)
    add_bullets(
        doc,
        [
            "Application UI: Teacher attempts list after instructor launch.",
            "SQL via Docker: docker exec -it db-db-1 psql -U edvidura -d edvidura",
            "GUI tools: Connect to host 127.0.0.1, port 5433, database edvidura.",
        ],
    )
    add_para(
        doc,
        "Primary tables: tenants, lti_platforms, launch_events, quiz_attempts, lti_launch_snapshots, quiz_session_tokens.",
        justify=True,
    )

    add_heading_numbered(doc, "9. Explicitly out of scope (not completed — deferred)")
    add_table(
        doc,
        ["Area", "Status"],
        [
            ("AI tutoring / generation features", "Deferred"),
            ("XR / simulation modules", "Deferred"),
            ("Federation / Keycloak enterprise IdP", "Deferred"),
            ("Air-gapped defence enclave deployment", "Deferred"),
            ("Full production observability / HA / multi-region", "Not in spike"),
        ],
        col_widths=[4.5, 2.2],
    )

    add_heading_numbered(doc, "10. Known limitations and residual risks")
    add_bullets(
        doc,
        [
            "AGS success depends on correct Moodle external-tool and activity settings.",
            "Local development uses in-memory caches plus Postgres snapshots; production should use a shared cache or durable session store.",
            "Stage 0 contract pack still requires formal leadership sign-off where drafted.",
            "Onboarding remains API-first; a minimal admin UI is recommended for demo operators.",
        ],
    )

    add_heading_numbered(doc, "11. Conclusion")
    add_para(
        doc,
        "The EdVidura LTI Hello spike has achieved its stated goal for Slice A: a multi-tenant, "
        "Moodle-launched quiz with tenant-isolated persistence and optional grade return to the LMS. "
        "The technical foundation (LTI 1.3, PostgreSQL RLS, onboarding APIs, isolation tests) is in place "
        "for subsequent hardening and later product slices.",
        justify=True,
    )
    add_para(
        doc,
        "Overall status of scoped spike work: COMPLETE for Slice A delivery criteria, subject to Moodle "
        "AGS configuration for gradebook demonstration and routine operational polish.",
        justify=True,
        bold=True,
    )

    doc.add_paragraph()
    sig = doc.add_table(rows=1, cols=2)
    left, right = sig.rows[0].cells
    for cell_, title in ((left, "Prepared by:"), (right, "Reviewed / acknowledged by:")):
        cell_.text = ""
        run = cell_.paragraphs[0].add_run(title)
        set_run_font(run, size=11, bold=True)
        for line in (
            "",
            "Name / Role: _______________________________",
            "",
            "Date: _______________    Signature: _______________",
        ):
            p = cell_.add_paragraph()
            run = p.add_run(line)
            set_run_font(run, size=11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(
        "EdVidura LTI Hello — Formal Progress Report (Completed Work) · 11 August 2026 · "
        "Black & white print version."
    )
    set_run_font(run, size=9)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
